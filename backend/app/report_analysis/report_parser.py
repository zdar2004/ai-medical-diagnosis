"""
report_parser.py
================

Medical Report Parser.

This module is responsible for extracting raw text from uploaded medical
report files before the cleaning and analysis stages.

Supported formats
-----------------
- PDF (.pdf)
- DOCX (.docx)
- TXT (.txt)

Responsibilities
----------------
- Validate uploaded files
- Detect supported formats
- Extract raw text
- Return plain text only

This module NEVER:

- Cleans text
- Detects diseases
- Extracts laboratory values
- Uses LLMs
- Performs interpretation
"""

from __future__ import annotations

import io
from pathlib import Path
from typing import BinaryIO, Optional, Union

import pdfplumber
from docx import Document

from app.risk_assessment.utils.logging_utils import get_logger

logger = get_logger(__name__)


SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".txt",
}

DEFAULT_ENCODING = "utf-8"

MAX_FILE_SIZE_MB = 20


class MedicalReportParser:
    """
    Parse uploaded medical reports.

    This class extracts plain text from supported medical report files.

    Supported formats

    - PDF
    - DOCX
    - TXT

    The parser performs only extraction.

    No cleaning or interpretation is performed here.
    """

    def __init__(
        self,
        max_file_size_mb: int = MAX_FILE_SIZE_MB,
    ) -> None:
        """
        Initialize parser.

        Args:
            max_file_size_mb:
                Maximum allowed uploaded file size.
        """

        self.max_file_size_mb = max_file_size_mb

        logger.info(
            "MedicalReportParser initialized."
        )
        
        # ---------------------------------------------------------
    # Validation
    # ---------------------------------------------------------

    def _validate_extension(
        self,
        file_path: Path,
    ) -> None:
        """
        Validate supported extension.

        Args:
            file_path:
                Report file path.

        Raises:
            ValueError:
                Unsupported extension.
        """

        extension = file_path.suffix.lower()

        if extension not in SUPPORTED_EXTENSIONS:

            logger.error(
                "Unsupported report type: %s",
                extension,
            )

            raise ValueError(
                f"Unsupported file type: {extension}"
            )

    def _validate_exists(
        self,
        file_path: Path,
    ) -> None:
        """
        Ensure file exists.

        Args:
            file_path:
                Report path.

        Raises:
            FileNotFoundError
        """

        if not file_path.exists():

            logger.error(
                "Report not found: %s",
                file_path,
            )

            raise FileNotFoundError(file_path)

    def _validate_size(
        self,
        file_path: Path,
    ) -> None:
        """
        Validate maximum upload size.

        Args:
            file_path:
                Report path.

        Raises:
            ValueError
        """

        size_mb = file_path.stat().st_size / (
            1024 * 1024
        )

        if size_mb > self.max_file_size_mb:

            logger.error(
                "Report exceeds maximum size."
            )

            raise ValueError(
                f"Maximum file size is "
                f"{self.max_file_size_mb} MB."
            )
            
    def _validate_file(
        self,
        file_path: Union[str, Path],
    ) -> Path:
        """
        Perform complete validation.

        Args:
            file_path:
                Uploaded report path.

        Returns:
            Validated Path object.
        """

        path = Path(file_path)

        self._validate_exists(path)

        self._validate_extension(path)

        self._validate_size(path)

        return path
    
        # ---------------------------------------------------------
    # PDF Extraction
    # ---------------------------------------------------------

    def _extract_pdf(
        self,
        file_path: Path,
    ) -> str:
        """
        Extract text from a PDF report.

        Each page is processed independently. Empty pages are skipped.
        The extracted text is returned as a single string separated by
        blank lines.

        Args:
            file_path:
                Path to the PDF report.

        Returns:
            Extracted raw text.

        Raises:
            RuntimeError:
                If PDF extraction fails.
        """

        logger.info(
            "Starting PDF extraction: %s",
            file_path.name,
        )

        pages: list[str] = []

        try:

            with pdfplumber.open(file_path) as pdf:

                total_pages = len(pdf.pages)

                logger.info(
                    "PDF contains %d pages.",
                    total_pages,
                )

                for page_number, page in enumerate(
                    pdf.pages,
                    start=1,
                ):

                    logger.debug(
                        "Extracting page %d.",
                        page_number,
                    )

                    page_text = page.extract_text()

                    if page_text is None:
                        logger.warning(
                            "Page %d contains no text.",
                            page_number,
                        )
                        continue

                    page_text = page_text.strip()

                    if not page_text:
                        logger.warning(
                            "Page %d is empty.",
                            page_number,
                        )
                        continue

                    pages.append(page_text)

        except Exception as error:

            logger.exception(
                "PDF extraction failed: %s",
                error,
            )

            raise RuntimeError(
                "Unable to extract text from PDF."
            ) from error

        if not pages:

            logger.warning(
                "No readable text found inside PDF."
            )

            return ""

        extracted_text = "\n\n".join(pages)

        logger.info(
            "PDF extraction completed successfully."
        )

        logger.debug(
            "Extracted %d characters.",
            len(extracted_text),
        )

        return extracted_text
    
    def _extract_pdf_from_bytes(
        self,
        file_data: bytes,
    ) -> str:
        """
        Extract text from PDF bytes.

        This helper allows future FastAPI UploadFile support
        without saving the PDF first.

        Args:
            file_data:
                Raw PDF bytes.

        Returns:
            Extracted text.

        Raises:
            RuntimeError:
                If extraction fails.
        """

        logger.info(
            "Starting in-memory PDF extraction."
        )

        pages: list[str] = []

        try:

            with pdfplumber.open(
                io.BytesIO(file_data)
            ) as pdf:

                for page in pdf.pages:

                    text = page.extract_text()

                    if text:

                        text = text.strip()

                        if text:

                            pages.append(text)

        except Exception as error:

            logger.exception(
                "In-memory PDF extraction failed."
            )

            raise RuntimeError(
                "Unable to parse PDF bytes."
            ) from error

        return "\n\n".join(pages)
    
        # ---------------------------------------------------------
    # DOCX Extraction
    # ---------------------------------------------------------

    def _extract_docx(
        self,
        file_path: Path,
    ) -> str:
        """
        Extract text from a DOCX medical report.

        Args:
            file_path:
                Path to the DOCX file.

        Returns:
            Plain extracted text.

        Raises:
            RuntimeError:
                If extraction fails.
        """

        logger.info(
            "Starting DOCX extraction: %s",
            file_path.name,
        )

        try:

            document = Document(file_path)

            paragraphs: list[str] = []

            for paragraph in document.paragraphs:

                text = paragraph.text.strip()

                if text:

                    paragraphs.append(text)

            extracted_text = "\n".join(paragraphs)

            logger.info(
                "DOCX extraction completed successfully."
            )

            logger.debug(
                "Extracted %d characters.",
                len(extracted_text),
            )

            return extracted_text

        except Exception as error:

            logger.exception(
                "DOCX extraction failed: %s",
                error,
            )

            raise RuntimeError(
                "Unable to extract text from DOCX."
            ) from error
            
        # ---------------------------------------------------------
    # TXT Extraction
    # ---------------------------------------------------------

    def _extract_txt(
        self,
        file_path: Path,
        encoding: str = DEFAULT_ENCODING,
    ) -> str:
        """
        Extract text from a TXT report.

        Args:
            file_path:
                Path to the text file.

            encoding:
                Text encoding.

        Returns:
            Plain text.

        Raises:
            RuntimeError:
                If extraction fails.
        """

        logger.info(
            "Starting TXT extraction: %s",
            file_path.name,
        )

        try:

            extracted_text = file_path.read_text(
                encoding=encoding,
                errors="ignore",
            )

            logger.info(
                "TXT extraction completed successfully."
            )

            logger.debug(
                "Extracted %d characters.",
                len(extracted_text),
            )

            return extracted_text

        except Exception as error:

            logger.exception(
                "TXT extraction failed: %s",
                error,
            )

            raise RuntimeError(
                "Unable to extract text from TXT."
            ) from error
            
        # ---------------------------------------------------------
    # Generic Extraction Dispatcher
    # ---------------------------------------------------------

    def _extract_text(
        self,
        file_path: Path,
    ) -> str:
        """
        Dispatch extraction based on file extension.

        Args:
            file_path:
                Validated report path.

        Returns:
            Raw extracted text.

        Raises:
            ValueError:
                Unsupported file extension.
        """

        extension = file_path.suffix.lower()

        logger.info(
            "Detected report type: %s",
            extension,
        )

        if extension == ".pdf":

            return self._extract_pdf(file_path)

        if extension == ".docx":

            return self._extract_docx(file_path)

        if extension == ".txt":

            return self._extract_txt(file_path)

        raise ValueError(
            f"Unsupported report type: {extension}"
        )
        
        # ---------------------------------------------------------
    # Public API
    # ---------------------------------------------------------

    def parse(
        self,
        file_path: Union[str, Path],
    ) -> str:
        """
        Parse a medical report file and extract raw text.

        This is the main public interface of the parser.

        Workflow:

        1. Validate file
        2. Detect file type
        3. Extract text
        4. Return raw extracted content

        Args:
            file_path:
                Path to uploaded medical report.

        Returns:
            Extracted raw report text.

        Raises:
            FileNotFoundError:
                If file does not exist.

            ValueError:
                If file type is unsupported.

            RuntimeError:
                If extraction fails.
        """

        logger.info(
            "Medical report parsing started."
        )

        validated_path = self._validate_file(
            file_path
        )

        extracted_text = self._extract_text(
            validated_path
        )

        if not extracted_text.strip():

            logger.warning(
                "Parser completed but no text was extracted."
            )

        logger.info(
            "Medical report parsing completed."
        )

        return extracted_text
    
    def parse_bytes(
        self,
        file_data: bytes,
        filename: str,
    ) -> str:
        """
        Parse report from uploaded bytes.

        This method is designed for FastAPI UploadFile.

        Args:
            file_data:
                Raw uploaded file bytes.

            filename:
                Original file name.

        Returns:
            Extracted raw text.

        Raises:
            ValueError:
                Unsupported extension.
        """

        logger.info(
            "Parsing report from bytes: %s",
            filename,
        )

        extension = Path(
            filename
        ).suffix.lower()

        if extension not in SUPPORTED_EXTENSIONS:

            raise ValueError(
                f"Unsupported file type: {extension}"
            )

        try:

            if extension == ".pdf":

                return self._extract_pdf_from_bytes(
                    file_data
                )

            if extension == ".txt":

                return file_data.decode(
                    DEFAULT_ENCODING,
                    errors="ignore",
                )

            if extension == ".docx":

                document = Document(
                    io.BytesIO(file_data)
                )

                paragraphs = []

                for paragraph in document.paragraphs:

                    text = paragraph.text.strip()

                    if text:

                        paragraphs.append(text)

                return "\n".join(paragraphs)

        except Exception as error:

            logger.exception(
                "Byte parsing failed."
            )

            raise RuntimeError(
                "Unable to parse uploaded file."
            ) from error

        return ""
    
    